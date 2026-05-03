"""CV parsing pipeline.

Stages:
  1. File → raw text (pdfplumber for PDF, python-docx for DOCX).
  2. Raw text → structured fields via heuristic section detection.

Design notes:
  * Heuristic-only — no external LLM calls. Robust to messy formatting.
  * Section headers, bullet splitting, contact extraction live in
    `app.utils.text_cleaning` so they can be reused and unit-tested.
  * Parser never raises on malformed input: empty fields stay empty.
"""
from __future__ import annotations

import io
import re
from dataclasses import asdict, dataclass, field

# pdfplumber / python-docx are imported lazily inside the extractors so the
# pure text-parsing path stays usable in environments where they aren't
# installed (e.g. running unit tests on the heuristic logic).

from app.utils.text_cleaning import (
    clean_text,
    extract_contacts,
    is_section_header,
    remove_contact_lines,
    split_csv_like,
    split_entries,
    SECTION_HEADERS,
)


# ------- Text extraction -------

def extract_text_from_pdf(data: bytes) -> str:
    """Extract text from a PDF byte buffer using pdfplumber."""
    import pdfplumber  # local import — keeps heuristic path dep-free.

    out: list[str] = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            out.append(page_text)
    return clean_text("\n".join(out))


def extract_text_from_docx(data: bytes) -> str:
    """Extract text from a DOCX byte buffer using python-docx."""
    from docx import Document  # local import — see note above.

    doc = Document(io.BytesIO(data))
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text:
            parts.append(para.text)
    # Include table cells too — many CVs use table-based layouts.
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    return clean_text("\n".join(parts))


def extract_text_from_txt(data: bytes) -> str:
    """Decode a UTF-8 (BOM-tolerant) TXT byte buffer."""
    for enc in ("utf-8-sig", "utf-8", "latin-1"):
        try:
            return clean_text(data.decode(enc))
        except UnicodeDecodeError:
            continue
    return clean_text(data.decode("utf-8", errors="replace"))


def extract_text(data: bytes, ext: str) -> str:
    """Dispatch to the right extractor by extension. Supports .pdf / .docx / .txt."""
    ext = ext.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(data)
    if ext == ".docx":
        return extract_text_from_docx(data)
    if ext == ".txt":
        return extract_text_from_txt(data)
    raise ValueError(f"Unsupported extension: {ext}")


# ------- Structured parsing -------

# Words that disqualify a line from being a candidate name.
_NAME_BLOCKLIST = {
    "curriculum", "vitae", "resume", "cv", "profile", "summary",
    "objective", "address", "phone", "email", "contact",
}
_NAME_LINE_RE = re.compile(r"^[A-Za-zÀ-ÿ\.\-' ]{2,60}$")


@dataclass
class ParsedCV:
    """Container returned by `parse_cv_text`. All fields default to empty."""
    name: str = ""
    summary: str = ""
    skills: list[str] = field(default_factory=list)
    education: list[str] = field(default_factory=list)
    experience: list[str] = field(default_factory=list)
    projects: list[str] = field(default_factory=list)
    certifications: list[str] = field(default_factory=list)
    languages: list[str] = field(default_factory=list)

    # Contact info.
    email: str = ""
    phone: str = ""
    linkedin: str = ""
    github: str = ""
    portfolio: str = ""

    def to_dict(self) -> dict:
        return asdict(self)


def _guess_name(lines: list[str]) -> str:
    """Pick the first plausible name line at the top of the CV."""
    candidates = remove_contact_lines(lines[:10])
    for ln in candidates:
        stripped = ln.strip()
        if not stripped or len(stripped) > 60:
            continue
        low = stripped.lower()
        if any(b in low for b in _NAME_BLOCKLIST):
            continue
        if not _NAME_LINE_RE.match(stripped):
            continue
        words = stripped.split()
        if not (1 < len(words) <= 5):
            continue
        # Title-case-ish: at least one capitalized word.
        if not any(w[:1].isupper() for w in words):
            continue
        return stripped
    return ""


def _build_sections(lines: list[str]) -> tuple[dict[str, list[str]], list[str]]:
    """Walk lines, grouping content under the most-recently-seen section.

    Returns:
        (sections_map, pre_section_lines) where pre_section_lines hold lines
        before the first recognised header (used for name / summary fallback).
    """
    sections: dict[str, list[str]] = {k: [] for k in SECTION_HEADERS}
    pre: list[str] = []
    current: str | None = None
    for ln in lines:
        header = is_section_header(ln)
        if header is not None:
            current = header
            continue
        if current is None:
            pre.append(ln)
        else:
            sections[current].append(ln)
    return sections, pre


def _summary_from(sections: dict[str, list[str]], pre: list[str]) -> str:
    """Prefer an explicit Summary section, else use the top of the CV."""
    block = "\n".join(sections["summary"]).strip()
    if block:
        # First entry only — multi-paragraph summaries get noisy.
        entries = split_entries(block)
        if entries:
            return entries[0][:1000]
        return block[:1000]
    # Fallback: first prose lines after stripping contact junk.
    body = [ln for ln in remove_contact_lines(pre) if ln.strip()]
    if not body:
        return ""
    # Skip the first line if it was the candidate's name.
    name_guess = _guess_name(pre)
    if name_guess and body and body[0].strip() == name_guess:
        body = body[1:]
    soft = " ".join(body[:5]).strip()
    return soft[:1000]


def _languages_from(sections: dict[str, list[str]]) -> list[str]:
    """Languages section is short — split CSV-like."""
    block = "\n".join(sections["languages"]).strip()
    if not block:
        return []
    items = split_csv_like(block, max_item_len=60)
    if items:
        return items
    # Some CVs put languages on multi-line bullets ("English – Fluent").
    return split_entries(block)


def parse_cv_text(text: str) -> ParsedCV:
    """Heuristic section parser. Returns a `ParsedCV` dataclass.

    Never raises on malformed input. Missing fields are returned as empty
    strings or lists.
    """
    text = clean_text(text or "")
    if not text:
        return ParsedCV()

    lines = text.split("\n")
    parsed = ParsedCV()

    # 1. Contact info comes from the entire document (often spread over lines).
    contacts = extract_contacts(text)
    parsed.email = contacts["email"]            # type: ignore[assignment]
    parsed.phone = contacts["phone"]            # type: ignore[assignment]
    parsed.linkedin = contacts["linkedin"]      # type: ignore[assignment]
    parsed.github = contacts["github"]          # type: ignore[assignment]
    parsed.portfolio = contacts["portfolio"]    # type: ignore[assignment]

    # 2. Name from the top, ignoring contact lines.
    parsed.name = _guess_name(lines)

    # 3. Section walk.
    sections, pre = _build_sections(lines)

    parsed.summary = _summary_from(sections, pre)
    parsed.skills = split_csv_like("\n".join(sections["skills"]))
    parsed.education = split_entries("\n".join(sections["education"]))
    parsed.experience = split_entries("\n".join(sections["experience"]))
    parsed.projects = split_entries("\n".join(sections["projects"]))
    parsed.certifications = split_entries("\n".join(sections["certifications"]))
    parsed.languages = _languages_from(sections)

    return parsed


# ------- Demo / smoke test -------

SAMPLE_CV = """\
Jane Doe
San Francisco, CA | jane.doe@example.com | +1 (415) 555-0199
linkedin.com/in/janedoe  |  github.com/janedoe  |  janedoe.dev

PROFESSIONAL SUMMARY
Backend engineer with 6 years building distributed systems in Python and Go.
Passionate about developer tooling, observability, and clean APIs.

TECHNICAL SKILLS
Python, Go, FastAPI, PostgreSQL, Redis, Kafka, Docker, Kubernetes, AWS, Terraform

WORK EXPERIENCE
Senior Backend Engineer — Acme Corp (2021 - Present)
- Designed a multi-tenant billing service handling 5M events/day.
- Cut p99 latency 40% by introducing async batching and read replicas.

Backend Engineer — Globex (2018 - 2021)
- Built the core payments API in Python/FastAPI.
- Led migration from monolith to 6 microservices on Kubernetes.

EDUCATION
B.Sc. Computer Science — University of California, Berkeley (2014 - 2018)

PROJECTS
- OpenObserve: open-source observability dashboard, 1.2k GitHub stars.
- pg-tuner: CLI that recommends Postgres config based on workload.

CERTIFICATIONS
- AWS Certified Solutions Architect — Associate (2022)
- Certified Kubernetes Administrator (CKA), 2021

LANGUAGES
English (Native), Spanish (Conversational), German (Basic)
"""


def run_sample() -> dict:
    """Run the parser on `SAMPLE_CV` and return the parsed dict.

    Useful as a quick smoke test:
        python -m app.services.cv_parser
    """
    return parse_cv_text(SAMPLE_CV).to_dict()


if __name__ == "__main__":  # pragma: no cover
    import json
    print(json.dumps(run_sample(), indent=2))
